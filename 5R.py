import random

from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE
import os

docPath = "C:\\Users\\Li Chung Ho\\Downloads\\Progress photo report\\ZHEC Progress photo report May - 5R (Sample).docx"
imgDirPath = "C:\\Users\\Li Chung Ho\\Downloads\\May 2023"
outputPath = "C:\\Users\\Li Chung Ho\\Downloads\\Output.docx"


def bold_underline(runner):
    runner.bold = True
    runner.underline = True


# This function return the Month MM of image name
def getMonth(s):
    return int(s.rstrip(".JPGjpg").split()[0].split("-")[0])


# This function return the Date DD of image name
def getDate(s):
    return int(s.rstrip(".JPGjpg").split()[0].split("-")[1])


def add0_formatting(x):
    int(x)
    return str(x) if x >= 10 else "0" + str(x)

# This function reformat img name to "DD/MM/YYYY"
def formatDate(imgName):
    date = getDate(imgName)
    month = getMonth(imgName)
    result = ""
    result += add0_formatting(date)
    result += "/"
    result += add0_formatting(month)
    result += "/2023"
    return result


# in the format of "HH:MM"
def generateRandomTime():
    minute = int(random.random()*60)
    hour = int(random.random()*(17-8)+8)
    # hour within 8am to 5pm
    result = ""
    result += add0_formatting(hour)
    result += ":"
    result += add0_formatting(minute)
    return result


doc = Document(docPath)

# set Style: font name and size
doc.styles.add_style("Big", WD_STYLE_TYPE.PARAGRAPH)
font = doc.styles["Big"].font
font.size = Pt(14)
font.name = "Arial"

doc.styles.add_style("Small", WD_STYLE_TYPE.PARAGRAPH)
font = doc.styles["Small"].font
font.size = Pt(10)
font.name = "Arial"

# first check if image "DD-MM.jpg" or "DD-MM (x).jpg" exist
# all image path are stored into a list temp
# after sorting, put it into another list imgPathList
temp = os.listdir(imgDirPath)
imgPathList = sorted(temp, key=getDate)
print(imgPathList)

for i in range(36):
    # Image replacement

    currentCell = doc.tables[2*i].rows[0].cells[0]
    currentCell._element.clear_content()
    p = currentCell.add_paragraph()
    p.add_run().add_picture(imgDirPath + "\\" + imgPathList[i], height=Inches(5.51))
    p.alignment = 1

    currentCell = doc.tables[2*i].rows[1].cells[0]
    currentCell._element.clear_content()
    text = str(i+1) if i+1>=10 else "0"+str(i+1)
    text = "PHOTO " + text + ":"
    p = currentCell.add_paragraph(text)
    p.style = doc.styles["Big"]
    p.alignment = 2

    currentCell = doc.tables[2*i].rows[1].cells[1]
    currentCell._element.clear_content()
    currentCell.add_paragraph("LOCATION – ").style = doc.styles["Big"]

    currentCell = doc.tables[2*i].rows[2].cells[1]
    currentCell._element.clear_content()
    currentCell.add_paragraph("PHOTOGRAPH TITLE – ").style = doc.styles["Big"]

    # Text replacement
    currentCell = doc.tables[2*i+1].rows[0].cells[0]
    currentCell._element.clear_content()
    p = currentCell.add_paragraph("CONTRACT NO. ")
    p.style = doc.styles["Small"]
    bold_underline(p.add_run("CC/2022/05/094"))
    p.add_run("       " + str(i+1) + " OF 36\n\nLOCATION ")
    bold_underline(p.add_run("__"))
    p.add_run("\n\nTITLE   ")
    bold_underline(p.add_run("__"))
    p.add_run("\n\nDATE PHOTOGRAPH TAKEN ")
    bold_underline(p.add_run(formatDate(imgPathList[i])))
    p.add_run("\n\nTIME ")
    bold_underline(p.add_run(generateRandomTime()))
    p.add_run("\n\nTHE CONTRACTOR'S SIGNATURE __________________")
    p.add_run("\n\nTHE SUPERVISOR'S SIGNATURE __________________")
    p.add_run("\n\nCERTIFYING THE ABOVE INFORMATION IS CORRECT.")
    p.add_run("\n\nPHOTOGRAPH TAKEN BY ")
    bold_underline(p.add_run("ZHEN HUA ENGINEERING CO,. LTD"))

doc.save(outputPath)
os.startfile(outputPath)

