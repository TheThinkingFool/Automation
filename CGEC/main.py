"""
Standard text styles in the DOCX file:
    - Arial, 12pt, left alignment
    - others: <bu>/<bu> stands for bold and underline
    - others: $var$ stands for variable

variables:
    - $location$: location of the photo
    - $title$: title of the photo, e.g. Aerial View, General View of ...
    - $ref_no$: reference number of the photo, range from 01 to 36
    - $date_time$: date and time of the photo, e.g. 31/01/2024 13:59
    - $weather_condition$: weather condition of the photo, either "FINE" or "CLOUDY"

Table Structure and their styles:
    - [Photos], really photos in .png format
    - ['PHOTO 4R$ref_no$: LOCATION – $location$'], Style: Calibri, 12pt, centered alignment
    - ['PHOTOGRAPH TITLE – $title$'], Style: Calibri, 12pt, centered alignment
    - ['']
    - ['CONTRACT NO. <bu>3/WSD/20</bu>\t$ref_no$ OF 36']
    - ['REFERENCE NUMBER <bu>4R$ref_no$</bu>']
    - ['LOCATION\t<bu>$location$</bu>']
    - ['WEATHER CONDITIONS <bu>$weather_condition$</bu>']
    - ['TITLE\t<bu>$title$</bu>']
    - ['DATE PHOTOGRAPH TAKEN ON <bu>$date_time$</bu>']
    - ["THE CONTRACTOR'S SIGNATURE _________________"]
    - ["THE SUPERVISOR’S SIGNATURE _________________"]
    - ['CERTIFYING THE ABOVE INFORMATION IS CORRECT.']
    - ['PHOTOGRAPH TAKEN BY <bu>CHAN TSZ KIN</bu>']

Plan:
    - extract the variables from the images name of the image folder
        image name format: $ref_no$_$location$_$title$_$weather_condition$.jpg
        or $ref_no$_$location$_$title$.jpg, default weather condition is "FINE"
    - generate a random date and time for each photo, i.e. date = {random DD}/MM/YYYY, time = {random hh}: {random mm} from 09:00 to 16:59.
        ask the MM/YYYY of monthly progress photo from the user.
    - generate another image folder and copy the images from the original image folder,
        and stick the date and time on the bottom right of images,
        rename the images with the generated date and time. (format of $ref_no$_$location$_$title$_$date_time$_$weather_condition$.jpg)
    - replace the photos and variables with actual values and correct styles in the DOCX file.
    - save new DOCX file with the updated values.
"""

import sys
import tkinter as tk
from tkinter import filedialog
import extract_image_variables as eiv
import generate_date_time as gdt
import generate_images_with_datetime as gidt
import os
import docx
from docx import Document
from docx.shared import Pt
from docx.shared import Cm
from PIL import Image

def open_folder_explorer(str="Select a Folder"):
    root = tk.Tk()
    root.withdraw()  # Hide the root window
    folder_path = filedialog.askdirectory(title=str)
    return folder_path

def open_word_file_explorer(prompt="Select a word document (.doc or .docx)"):
    root = tk.Tk()
    root.withdraw()  # Hide the root window
    file_path = filedialog.askopenfilename(
        title=prompt,
        filetypes=[("Word Documents", "*.doc *.docx")]
    )
    return file_path

# part 1 in extract_image_variables.py, imported as eiv
image_folder = open_folder_explorer("Select a folder containing Photos.")
if image_folder:
    print(f"Selected folder: {image_folder}")
else:
    sys.exit("No folder selected.")
image_variables = eiv.extract_image_variables(image_folder)
print("Extracted Image Variables:")
for var in image_variables:
    print(var)


# part 2 in generate_date_time.py, imported as gdt
month_year = gdt.get_month_year()
num_photos = len(image_variables) # image_variables extracted from part 1
date_times = gdt.generate_random_date_time(month_year, num_photos)
print("Generated Date and Time for each photo:")
for dt in date_times:
    print(dt)

# part 3 in generate_images_with_datetime.py, imported as gidt

print(image_folder)
print(image_variables)
output_folder = open_folder_explorer("Select a folder to save the images with date and time.")
output_images_path = gidt.process_images_with_datetime(image_folder, output_folder, image_variables, date_times)

# part 4 here
'''
Part 4: Replace Photos and Variables in DOCX File
    - replace the photos and variables with actual values and correct styles in the DOCX file.
    - save new DOCX file with the updated values.
'''

def bold_underline(runner):
    runner.bold = True
    runner.underline = True

def replace_photos_and_variables(docx_path, output_path, image_folder, image_variables, date_times):
    doc = Document(docx_path)

    # Replace photos in the DOCX file
    for i, image_var in enumerate(image_variables):
        ref_no, location, title, weather_condition = image_var.values()
        current_no = int(ref_no) - 1

        image_path = output_images_path[i]
        currentCell = doc.tables[current_no].rows[0].cells[0]
        currentCell._element.clear_content()
        p = currentCell.add_paragraph()
        with Image.open(image_path) as img:
            width, height = img.size
            aspect_ratio = width / height

            # Calculate the new dimensions
            if width / 15.5 > height / 10.5:
                new_width = Cm(15.5)
                new_height = Cm(15.5 / aspect_ratio)
            else:
                new_height = Cm(10.5)
                new_width = Cm(10.5 * aspect_ratio)

        p.add_run().add_picture(image_path, width=new_width, height=new_height)
        p.alignment = 1

    # Replace variables in the DOCX file
    for i, image_var in enumerate(image_variables):
        ref_no, location, title, weather_condition = image_var.values()
        current_no = int(ref_no) - 1
        date_time = date_times[i]

        # second row, 'PHOTO 4R$ref_no$: LOCATION – $location$',
        # Style: Calibri, 12pt, centered alignment
        currentCell = doc.tables[current_no].rows[1].cells[0]
        currentCell._element.clear_content()
        p = currentCell.add_paragraph()
        run = p.add_run(f'PHOTO 4R{ref_no}: LOCATION – {location}')
        run.font.name = 'Calibri'
        run.font.size = Pt(12)
        p.alignment = 1

        # third row, 'PHOTOGRAPH TITLE – $title$',
        # Style: Calibri, 12pt, centered alignment
        currentCell = doc.tables[current_no].rows[2].cells[0]
        currentCell._element.clear_content()
        p = currentCell.add_paragraph()
        run = p.add_run(f'PHOTOGRAPH TITLE – {title}')
        run.font.name = 'Calibri'
        run.font.size = Pt(12)
        p.alignment = 1

        # fifth row, 'CONTRACT NO. <bu>3/WSD/20</bu>\t$ref_no$ OF 36'
        currentCell = doc.tables[current_no].rows[4].cells[0]
        currentCell._element.clear_content()
        p = currentCell.add_paragraph("CONTRACT NO. ")
        bold_underline(p.add_run("3/WSD/20"))
        p.add_run(f"\t{ref_no} OF 36")

        # sixth row, 'REFERENCE NUMBER <bu>4R$ref_no$</bu>'
        currentCell = doc.tables[current_no].rows[5].cells[0]
        currentCell._element.clear_content()
        p = currentCell.add_paragraph("REFERENCE NUMBER ")
        bold_underline(p.add_run(f"4R{ref_no}"))

        # seventh row, 'LOCATION\t<bu>$location$</bu>'
        currentCell = doc.tables[current_no].rows[6].cells[0]
        currentCell._element.clear_content()
        p = currentCell.add_paragraph("LOCATION\t")
        bold_underline(p.add_run(location))

        # eighth row, 'WEATHER CONDITIONS <bu>$weather_condition$</bu>'
        currentCell = doc.tables[current_no].rows[7].cells[0]
        currentCell._element.clear_content()
        p = currentCell.add_paragraph("WEATHER CONDITIONS ")
        bold_underline(p.add_run(weather_condition))

        # ninth row, 'TITLE\t<bu>$title$</bu>'
        currentCell = doc.tables[current_no].rows[8].cells[0]
        currentCell._element.clear_content()
        p = currentCell.add_paragraph("TITLE\t")
        bold_underline(p.add_run(title))

        # tenth row, 'DATE PHOTOGRAPH TAKEN ON <bu>$date_time$</bu>'
        currentCell = doc.tables[current_no].rows[9].cells[0]
        currentCell._element.clear_content()
        p = currentCell.add_paragraph("DATE PHOTOGRAPH TAKEN ON ")
        bold_underline(p.add_run(date_time))

        # eleventh row, "THE CONTRACTOR'S SIGNATURE _________________"
        currentCell = doc.tables[current_no].rows[10].cells[0]
        currentCell._element.clear_content()
        p = currentCell.add_paragraph("THE CONTRACTOR'S SIGNATURE _________________")

        # twelfth row, "THE SUPERVISOR’S SIGNATURE _________________"
        currentCell = doc.tables[current_no].rows[11].cells[0]
        currentCell._element.clear_content()
        p = currentCell.add_paragraph("THE SUPERVISOR’S SIGNATURE _________________")


    doc.save(output_path)


docx_path = open_word_file_explorer()  # Path to the template DOCX file
output_path = "output.docx"  # Path to save the updated DOCX file
image_folder = output_folder  # Path to the folder containing images with date and time
replace_photos_and_variables(docx_path, output_path, image_folder, image_variables, date_times)
print(f"Updated DOCX file saved to {output_path}")